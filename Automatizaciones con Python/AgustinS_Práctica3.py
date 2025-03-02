import PyPDF2
import docx
import os
from pdf2image import convert_from_path

def extraer_texto(pdf_path, paginas=None):
    texto_extraido = ""
    with open(pdf_path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        num_paginas = len(reader.pages)
        paginas = paginas if paginas else range(num_paginas)
        for i in paginas:
            if i < num_paginas:
                texto_extraido += reader.pages[i].extract_text() + "\n"
    return texto_extraido

def extraer_imagenes(pdf_path, output_folder="imagenes_extraidas"):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    images = convert_from_path(pdf_path)
    image_paths = []
    for i, img in enumerate(images):
        img_path = os.path.join(output_folder, f"pagina_{i + 1}.png")
        img.save(img_path, "PNG")
        image_paths.append(img_path)
    return image_paths

def guardar_en_word(texto, imagenes, output_path):
    doc = docx.Document()
    doc.add_paragraph(texto)
    for img in imagenes:
        doc.add_picture(img)
    doc.save(output_path)

if __name__ == "__main__":
    nombre_pdf = input("Ingrese el nombre del archivo PDF (incluyendo .pdf): ")
    if not os.path.exists(nombre_pdf):
        print("El archivo no existe. Verifique el nombre e intente nuevamente.")
    else:
        seleccion = input("¿Desea extraer texto de páginas específicas? (s/n): ")
        paginas = None
        if seleccion.lower() == "s":
            paginas = input("Ingrese los números de páginas separados por comas (empezando en 0): ")
            paginas = [int(x) for x in paginas.split(",")]
        
        texto_extraido = extraer_texto(nombre_pdf, paginas)
        imagenes_extraidas = extraer_imagenes(nombre_pdf)
        guardar_en_word(texto_extraido, imagenes_extraidas, "informe.docx")
        print("Proceso completado. Se ha guardado en 'informe.docx'.")
